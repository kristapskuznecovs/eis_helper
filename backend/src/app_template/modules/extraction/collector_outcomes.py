#!/usr/bin/env python3
"""LLM-only outcome extraction for procurement reports."""

from __future__ import annotations

import html
import http.cookiejar
import io
import json
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Set
from urllib.request import HTTPCookieProcessor, Request, build_opener
from urllib.parse import urlencode

from .collector_classes import OutcomeLLMExtractor, paced_opener_open
from .document_extractor import get_archive_content_as_bytes
from .utils import extract_js_array, fetch_html, is_captcha_page, normalize_text


def parse_optional_float(value: Any) -> Optional[float]:
    """Parse a value to float, handling None, strings with spaces/commas."""
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().replace(" ", "")
    if not text:
        return None
    text = text.replace(",", ".")
    try:
        return float(text)
    except ValueError:
        return None


def trim_company_name_noise(name: str) -> str:
    """Remove trailing noise from company names."""
    patterns = [
        r"(.*?)\s*\(?rg\b.*",
        r"(.*?)\s*\(?reģ.*",
        r"(.*?)\s*\d{11}.*",
    ]
    for pattern in patterns:
        match = re.match(pattern, name, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return name


def canonical_company_name(text: str) -> str:
    """Canonicalize company name for comparison."""
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text


def company_name_core(name: str) -> str:
    """Extract core part of company name for matching."""
    normalized = normalize_text(name)
    form_pattern = r'\b(sia|as|ps|ik|zs|ab|piegadataju apvieniba)\b'
    core = re.sub(form_pattern, '', normalized, flags=re.IGNORECASE)
    core = re.sub(r'\s+', ' ', core).strip()
    return core


def is_plausible_company_name(name: str) -> bool:
    """Check if a name looks like a real company name."""
    if not name or len(name) < 3:
        return False
    normalized = normalize_text(name)
    if len(normalized) < 3:
        return False

    letter_count = sum(c.isalpha() for c in name)
    if letter_count < 2:
        return False

    false_positives = [
        "zemaka cena", "ligumcena", "eiro", "eur", "euro",
        "pretendenta nosaukums", "izveles pamatojums",
        "piedavajumu iesniegusi", "liguma slēgšanas tiesības"
    ]
    if normalized in false_positives:
        return False

    return True


def normalize_outcome_llm_result(parsed: Dict[str, Any]) -> Dict[str, Any]:
    """
    Normalize and clean LLM extraction result.

    Args:
        parsed: Raw LLM response dict

    Returns:
        Cleaned dict with procurement_status, winner_name, winner_registration_no,
        winner_price_eur, participants, confidence, notes
    """
    # Extract procurement status
    procurement_status = str(parsed.get("procurement_status") or "").strip().lower()
    if procurement_status not in {"completed", "cancelled", "terminated", "no_applications", "pre_consultation", "unknown"}:
        procurement_status = "unknown"

    participants_in = parsed.get("participants")
    participants: List[Dict[str, Any]] = []
    seen: Set[str] = set()

    if isinstance(participants_in, list):
        for item in participants_in:
            if not isinstance(item, dict):
                continue
            raw_name = str(item.get("name") or "").strip()
            if not raw_name:
                continue
            name = trim_company_name_noise(canonical_company_name(raw_name))
            if not is_plausible_company_name(name):
                continue
            key = normalize_text(name)
            if key in seen:
                continue
            seen.add(key)
            price = parse_optional_float(item.get("suggested_price_eur"))
            reg_no = str(item.get("registration_no") or "").strip() or None
            consortium_members = item.get("consortium_members")
            if consortium_members and isinstance(consortium_members, list):
                consortium_members = [str(m).strip() for m in consortium_members if m]
            else:
                consortium_members = None

            # Handle disqualification
            disqualified = bool(item.get("disqualified", False))
            disqualification_reason = str(item.get("disqualification_reason") or "").strip() or None

            participant_data = {
                "name": name,
                "registration_no": reg_no,
                "suggested_price_eur": price,
                "suggested_prices_eur": [price] if isinstance(price, float) else [],
                "disqualified": disqualified,
            }
            if disqualification_reason:
                participant_data["disqualification_reason"] = disqualification_reason
            if consortium_members:
                participant_data["consortium_members"] = consortium_members
            participants.append(participant_data)

    winner_raw = str(parsed.get("winner_name") or "").strip()
    winner_name = trim_company_name_noise(canonical_company_name(winner_raw)) if winner_raw else None
    winner_reg = str(parsed.get("winner_registration_no") or "").strip() or None
    winner_price = parse_optional_float(parsed.get("winner_price_eur"))
    confidence = str(parsed.get("confidence") or "").strip().lower() or None
    notes = str(parsed.get("notes") or "").strip() or None

    if winner_name:
        winner_core = company_name_core(winner_name)
        for participant in participants:
            participant_name = str(participant.get("name") or "")
            participant_core = company_name_core(participant_name)
            if winner_core and participant_core and (
                winner_core in participant_core or participant_core in winner_core
            ):
                winner_name = participant_name
                if not winner_reg:
                    winner_reg = str(participant.get("registration_no") or "").strip() or None
                if winner_price is None:
                    winner_price = parse_optional_float(participant.get("suggested_price_eur"))
                break

    # Extract new fields
    bid_deadline = str(parsed.get("bid_deadline") or "").strip() or None
    decision_date = str(parsed.get("decision_date") or "").strip() or None
    funding_source = str(parsed.get("funding_source") or "").strip() or None
    eu_project_reference = str(parsed.get("eu_project_reference") or "").strip() or None
    evaluation_method = str(parsed.get("evaluation_method") or "").strip() or None
    contract_scope_type = str(parsed.get("contract_scope_type") or "").strip() or None

    # Extract subcontractors
    subcontractors_in = parsed.get("subcontractors")
    subcontractors = None
    if subcontractors_in and isinstance(subcontractors_in, list):
        subcontractors = [str(s).strip() for s in subcontractors_in if s]
        if not subcontractors:
            subcontractors = None

    # Extract multi-lot fields
    is_multi_lot = parsed.get("is_multi_lot", False)
    lot_count = parsed.get("lot_count")
    if lot_count is not None:
        lot_count = int(lot_count) if lot_count else None

    return {
        "procurement_status": procurement_status,
        "winner_name": winner_name,
        "winner_registration_no": winner_reg,
        "winner_price_eur": winner_price,
        "participants": participants,
        "bid_deadline": bid_deadline,
        "decision_date": decision_date,
        "funding_source": funding_source,
        "eu_project_reference": eu_project_reference,
        "evaluation_method": evaluation_method,
        "contract_scope_type": contract_scope_type,
        "subcontractors": subcontractors,
        "is_multi_lot": is_multi_lot,
        "lot_count": lot_count,
        "confidence": confidence,
        "notes": notes,
    }


def is_finished_procurement_status(status_value: Any) -> bool:
    """Check if procurement status indicates completion."""
    status = normalize_text(str(status_value or ""))
    if not status:
        return False
    finished_markers = [
        "ligums noslegts",
        "noslegts",
        "contract concluded",
        "closed",
        "completed",
    ]
    return any(marker in status for marker in finished_markers)


def extract_text_from_doc_bytes(content: bytes) -> str:
    """
    Extract text from old .doc format (MS Word 97-2003).
    Uses textutil on macOS or antiword/catdoc on Linux.
    """
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp_doc:
            tmp_doc.write(content)
            tmp_doc_path = tmp_doc.name

        try:
            result = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", tmp_doc_path],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout

            result = subprocess.run(
                ["antiword", tmp_doc_path],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout

            result = subprocess.run(
                ["catdoc", tmp_doc_path],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout

            return ""
        finally:
            try:
                Path(tmp_doc_path).unlink()
            except OSError:
                pass
    except Exception:
        return ""


def extract_text_from_docx_bytes(content: bytes) -> str:
    """Extract text from DOCX file bytes (modern format)."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_from_pdf_bytes(content: bytes) -> str:
    """Extract text from PDF file bytes using pdfplumber."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            parts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            return "\n".join(parts)
    except Exception:
        return ""


def extract_text_from_pdf_bytes_ocr(content: bytes, max_pages: int = 25) -> str:
    """
    Extract text from PDF using OCR (tesseract).
    Used when direct text extraction returns empty or insufficient text.
    """
    try:
        import fitz  # PyMuPDF

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
            tmp_pdf.write(content)
            tmp_pdf_path = tmp_pdf.name

        try:
            doc = fitz.open(tmp_pdf_path)
            parts = []
            for page_num, page in enumerate(doc):
                if page_num >= max_pages:
                    break
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")

                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                    tmp_img.write(img_bytes)
                    tmp_img_path = tmp_img.name

                try:
                    result = subprocess.run(
                        ["tesseract", tmp_img_path, "stdout", "-l", "lav+eng"],
                        capture_output=True,
                        text=True,
                        timeout=30,
                    )
                    if result.returncode == 0:
                        parts.append(result.stdout)
                finally:
                    try:
                        Path(tmp_img_path).unlink()
                    except OSError:
                        pass

            doc.close()
            return "\n".join(parts)
        finally:
            try:
                Path(tmp_pdf_path).unlink()
            except OSError:
                pass
    except Exception:
        return ""


def convert_pdf_to_images_base64(content: bytes, max_pages: int = 10, dpi: int = 200) -> List[str]:
    """
    Convert PDF pages to base64-encoded PNG images for Vision API.

    Args:
        content: PDF file bytes
        max_pages: Maximum number of pages to convert (default 10)
        dpi: Image resolution (default 200)

    Returns:
        List of base64-encoded PNG images
    """
    try:
        import base64
        import fitz  # PyMuPDF

        images_b64 = []
        doc = fitz.open(stream=content, filetype="pdf")

        for page_num, page in enumerate(doc):
            if page_num >= max_pages:
                break

            pix = page.get_pixmap(dpi=dpi)
            img_bytes = pix.tobytes("png")
            img_b64 = base64.b64encode(img_bytes).decode("utf-8")
            images_b64.append(img_b64)

        doc.close()
        return images_b64
    except Exception:
        return []


def extract_text_from_report_bytes(filename: str, content: bytes) -> str:
    """
    Extract text from report document (PDF, DOCX, or archive).
    Uses OCR fallback if direct extraction returns insufficient text.

    Args:
        filename: Original filename (to determine type)
        content: File bytes

    Returns:
        Extracted text string
    """
    ext = Path(filename).suffix.lower()

    if ext in {".edoc", ".asice", ".zip", ".rar"}:
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp_path = Path(tmp.name)
            tmp.write(content)

        try:
            extracted_bytes = get_archive_content_as_bytes(tmp_path)
            if extracted_bytes:
                text = extract_text_from_pdf_bytes(extracted_bytes)
                if not text or len(text.strip()) < 100:
                    text = extract_text_from_doc_bytes(extracted_bytes)
                if not text or len(text.strip()) < 100:
                    text = extract_text_from_docx_bytes(extracted_bytes)
                if not text or len(text.strip()) < 100:
                    text = extract_text_from_pdf_bytes_ocr(extracted_bytes)
                return text
        finally:
            try:
                tmp_path.unlink()
            except OSError:
                pass
        return ""

    if ext == ".pdf":
        text = extract_text_from_pdf_bytes(content)
        if not text or len(text.strip()) < 100:
            text = extract_text_from_pdf_bytes_ocr(content)
        return text

    elif ext == ".doc":
        return extract_text_from_doc_bytes(content)

    elif ext == ".docx":
        return extract_text_from_docx_bytes(content)

    return ""


def select_final_report_document(actual_docs: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """
    Select the best outcome document from a list of documents.

    Priority:
    1) Final report (TypeCode "PRCFINSMR", "Noslēguma ziņojums")
    2) Opening meeting financial summary
       ("Atvēršanas sanāksmes finanšu piedāvājumu kopsavilkums" or EN equivalent)
    3) Generic report-like document as fallback
    """
    if not actual_docs:
        return None

    for doc in actual_docs:
        type_code = str(doc.get("TypeCode") or "")
        if type_code == "PRCFINSMR":
            return doc

        title = normalize_text(str(doc.get("Title") or ""))
        type_title = normalize_text(str(doc.get("TypeTitle") or ""))
        if "nosleguma zinojums" in title:
            return doc
        if "nosleguma zinojums" in type_title:
            return doc

    financial_summary_markers = (
        "atversanas sanaksmes finansu piedavajumu kopsavilkums",
        "finansu piedavajumu apkopoj",
        "opening meeting final financial proposal",
        "financial offers summary",
    )
    for doc in actual_docs:
        title = normalize_text(str(doc.get("Title") or ""))
        type_title = normalize_text(str(doc.get("TypeTitle") or ""))
        combined = f"{title} {type_title}"
        if any(marker in combined for marker in financial_summary_markers):
            return doc

    for doc in actual_docs:
        title = normalize_text(str(doc.get("Title") or ""))
        if "zinojums" in title or "report" in title:
            return doc

    return None


def extract_view_document_json_params(view_html: str) -> Dict[str, Any]:
    """Extract JSON parameters from ViewDocument page."""
    match = re.search(r'id="ViewDocumentModel_JsonParams"[^>]*value="([^"]*)"', view_html)
    if match:
        try:
            json_str = html.unescape(match.group(1))
            return json.loads(json_str)
        except json.JSONDecodeError:
            pass

    match = re.search(r"viewModel\s*=\s*(\{.*?\});", view_html, flags=re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass
    return {}


def extract_view_document_files(view_html: str) -> List[Dict[str, Any]]:
    """Extract file list from ViewDocument page."""
    match = re.search(r'var\s+ViewDocumentModel_Files_items\s*=\s*(\[.*?\]);', view_html, flags=re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass
    return []


def extract_final_report_outcome(
    project: Dict[str, Any],
    request_timeout_seconds: int,
    cookie_header: Optional[str],
    outcome_llm_extractor: Optional[OutcomeLLMExtractor],
) -> Dict[str, Any]:
    """
    Extract procurement outcome from final report using LLM only.

    Workflow:
    1. Fetch procurement page HTML
    2. Find and download final report document
    3. Extract text (with OCR fallback if needed)
    4. Send text to LLM for structured extraction
    5. Parse and normalize LLM response
    6. Save document if extraction successful

    Args:
        project: Project dict with procurement_url, procurement_id
        request_timeout_seconds: HTTP request timeout
        cookie_header: Optional cookie for authentication
        outcome_llm_extractor: LLM extractor instance (required)

    Returns:
        Dict with extracted outcome data
    """
    procurement_url = project.get("procurement_url", "")
    procurement_id = project.get("procurement_id", "unknown")

    base_result: Dict[str, Any] = {
        "procurement_winner": None,
        "procurement_winner_registration_no": None,
        "procurement_winner_suggested_price_eur": None,
        "procurement_winner_source": None,
        "procurement_participants_count": 0,
        "procurement_participants": [],
        "procurement_outcome_description": None,
    }

    if not procurement_url:
        base_result["procurement_outcome_description"] = "missing_procurement_url"
        return base_result

    try:
        page_html = fetch_html(procurement_url, request_timeout_seconds, cookie_header)
    except Exception as exc:
        base_result["procurement_outcome_description"] = f"page_fetch_error:{exc}"
        return base_result

    if is_captcha_page(page_html):
        base_result["procurement_outcome_description"] = "captcha_challenge"
        return base_result

    actual_docs = extract_js_array(page_html, "ActualDocuments_items")
    final_report_doc = select_final_report_document(actual_docs)

    if not final_report_doc:
        base_result["procurement_outcome_description"] = "final_report_not_found"
        return base_result

    file_id = final_report_doc.get("Id")
    file_name = final_report_doc.get("Title", "report") + ".bin"

    view_url = f"{procurement_url.rstrip('/')}/ViewDocument/{file_id}"
    try:
        view_html = fetch_html(view_url, request_timeout_seconds, cookie_header)
    except Exception as exc:
        base_result["procurement_outcome_description"] = f"view_page_error:{exc}"
        return base_result

    json_params = extract_view_document_json_params(view_html)
    document_data = json_params.get("Document") or {}
    file_name = document_data.get("FileName") or file_name

    base_url = procurement_url.split("/EKEIS/")[0] + "/EKEIS"
    download_params = {
        "Id": str(json_params.get("Id") or document_data.get("Id") or ""),
        "FileId": str(file_id),
        "DocumentLinkTypeCode": str(json_params.get("DocumentLinkTypeCode") or document_data.get("DocumentLinkTypeCode") or ""),
        "ParentObjectTypeCode": str(json_params.get("ParentObjectTypeCode") or document_data.get("ParentObjectTypeCode") or ""),
        "ParentId": str(json_params.get("ParentId") or document_data.get("ParentId") or ""),
        "ParentIdentifier": str(json_params.get("ParentIdentifier") or document_data.get("ParentIdentifier") or ""),
        "ProcurementIdentifier": str(json_params.get("ProcurementIdentifier") or document_data.get("ProcurementIdentifier") or procurement_id),
        "StageIdentifier": str(json_params.get("StageIdentifier") or document_data.get("StageIdentifier") or ""),
    }

    cookie_jar = http.cookiejar.CookieJar()
    opener = build_opener(HTTPCookieProcessor(cookie_jar))

    try:
        download_url = f"{base_url}/EKEIS/Document/DownloadDocumentFile?{urlencode(download_params)}"
        browser_ua = "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        download_request = Request(
            download_url,
            headers={"User-Agent": browser_ua, "Referer": procurement_url},
        )
        with paced_opener_open(opener, download_request, timeout=request_timeout_seconds) as response:
            report_bytes = response.read()
    except Exception as exc:
        base_result["procurement_outcome_description"] = f"report_download_error:{exc}"
        return base_result

    report_text = extract_text_from_report_bytes(file_name, report_bytes)

    if outcome_llm_extractor is None:
        base_result["procurement_outcome_description"] = "llm_extractor_not_configured"
        return base_result

    used_vision = False
    try:
        llm_parsed: Optional[Dict[str, Any]] = None
        text_is_insufficient = not report_text or len(report_text.strip()) < 100
        ext = Path(file_name).suffix.lower()

        if ext == ".pdf" and text_is_insufficient:
            images_b64 = convert_pdf_to_images_base64(report_bytes, max_pages=10)
            if images_b64:
                llm_parsed = outcome_llm_extractor.extract_from_images(
                    project=project,
                    file_name=file_name,
                    images_base64=images_b64,
                )
                used_vision = True

        if llm_parsed is None:
            if not report_text:
                base_result["procurement_outcome_description"] = "report_text_empty"
                return base_result
            llm_parsed = outcome_llm_extractor.extract(
                project=project,
                file_name=file_name,
                report_text=report_text,
            )

        llm_normalized = normalize_outcome_llm_result(llm_parsed)

        participants = llm_normalized.get("participants") or []
        winner_name = llm_normalized.get("winner_name")
        winner_reg_no = llm_normalized.get("winner_registration_no")
        winner_price = llm_normalized.get("winner_price_eur")
        confidence = llm_normalized.get("confidence") or "unknown"

    except Exception as exc:
        base_result["procurement_outcome_description"] = f"llm_extraction_error:{exc}"
        return base_result

    base_result["procurement_participants"] = participants
    base_result["procurement_participants_count"] = len(participants)
    base_result["procurement_winner"] = winner_name
    base_result["procurement_winner_registration_no"] = winner_reg_no
    base_result["procurement_winner_suggested_price_eur"] = winner_price
    base_result["procurement_winner_source"] = (
        "final_report_vision" if (winner_name and used_vision) else ("final_report_text" if winner_name else None)
    )
    llm_method = "vision" if used_vision else "text"
    base_result["procurement_outcome_description"] = (
        f"final_report:{final_report_doc.get('Title')}; file:{file_name}; "
        f"llm_method:{llm_method}; llm_confidence:{confidence}"
    )

    if winner_name or participants:
        try:
            save_dir = Path("data/construction_2023_2025/reports")
            save_dir.mkdir(parents=True, exist_ok=True)
            ext = Path(file_name).suffix.lower() or ".bin"

            if ext in {".edoc", ".asice", ".zip", ".rar"}:
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                    tmp_path = Path(tmp.name)
                    tmp.write(report_bytes)

                try:
                    extracted_bytes = get_archive_content_as_bytes(tmp_path)
                    if extracted_bytes:
                        extracted_ext = ".pdf"  # Default
                        if b"PK\x03\x04" in extracted_bytes[:4]:
                            extracted_ext = ".docx"

                        save_path = save_dir / f"{procurement_id}{extracted_ext}"
                        if not save_path.exists():
                            save_path.write_bytes(extracted_bytes)
                finally:
                    try:
                        tmp_path.unlink()
                    except OSError:
                        pass
            else:
                save_path = save_dir / f"{procurement_id}{ext}"
                if not save_path.exists():
                    save_path.write_bytes(report_bytes)
        except Exception:
            pass

    return base_result

